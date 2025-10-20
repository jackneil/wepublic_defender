"""
Document handlers for converting markdown legal documents to Word format.

Handles markdown-to-DOCX conversion with federal court formatting standards.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from typing import Optional, Dict, List
from pydantic import BaseModel, Field
import re
import json


class Party(BaseModel):
    """Represents a party in the case."""
    name: str
    type: str = "individual"  # individual, corporation, entity

class DocumentFormatConfig(BaseModel):
    """Configuration for federal court document formatting."""

    court_name: str = Field(
        default="IN THE UNITED STATES DISTRICT COURT",
        description="Name of the court"
    )
    court_district: str = Field(
        default="FOR THE [DISTRICT] OF [STATE]",
        description="Court district (must be provided at runtime)"
    )
    court_division: Optional[str] = Field(
        default=None,
        description="Court division (optional)"
    )
    case_number: str = Field(
        default="[CASE NUMBER]",
        description="Case number (must be provided at runtime)"
    )

    # Support multiple parties
    plaintiffs: List[Party] = Field(
        default=[Party(name="[PLAINTIFF NAME]", type="individual")],
        description="List of plaintiffs"
    )
    plaintiff_label: str = Field(
        default="Plaintiff",
        description="Plaintiff label (e.g., 'Plaintiff', 'Petitioner')"
    )
    defendants: List[Party] = Field(
        default=[Party(name="[DEFENDANT NAME]", type="individual")],
        description="List of defendants"
    )
    defendant_label: str = Field(
        default="Defendant",
        description="Defendant label (e.g., 'Defendant', 'Respondent')"
    )

    # Formatting options
    font_name: str = "Times New Roman"
    font_size: int = 12
    line_spacing: str = "double"  # "single" or "double"
    margin_inches: float = 1.0

    # Backward compatibility properties
    @property
    def plaintiff_name(self) -> str:
        """For backward compatibility - returns first plaintiff name."""
        return self.plaintiffs[0].name if self.plaintiffs else "[PLAINTIFF NAME]"

    @property
    def defendant_name(self) -> str:
        """For backward compatibility - returns first defendant name."""
        return self.defendants[0].name if self.defendants else "[DEFENDANT NAME]"

    @classmethod
    def from_case_config(cls, config_path: Optional[Path] = None) -> "DocumentFormatConfig":
        """Load configuration from case_config.json file.

        Args:
            config_path: Path to case_config.json. If None, looks for .wepublic_defender/case_config.json

        Returns:
            DocumentFormatConfig instance
        """
        if config_path is None:
            # Look for case config in standard location
            config_path = Path.cwd() / ".wepublic_defender" / "case_config.json"

        if not config_path.exists():
            # Return default config if file doesn't exist
            return cls()

        with open(config_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Map from case_config.json structure to DocumentFormatConfig
        config_dict = {}

        if 'court' in data:
            config_dict['court_name'] = data['court'].get('name', cls.model_fields['court_name'].default)
            config_dict['court_district'] = data['court'].get('district', cls.model_fields['court_district'].default)
            config_dict['court_division'] = data['court'].get('division')

        if 'parties' in data:
            if 'plaintiffs' in data['parties']:
                config_dict['plaintiffs'] = [Party(**p) for p in data['parties']['plaintiffs']]
            config_dict['plaintiff_label'] = data['parties'].get('plaintiff_label', 'Plaintiff')

            if 'defendants' in data['parties']:
                config_dict['defendants'] = [Party(**p) for p in data['parties']['defendants']]
            config_dict['defendant_label'] = data['parties'].get('defendant_label', 'Defendant')

        if 'case_number' in data:
            config_dict['case_number'] = data['case_number']

        if 'formatting' in data:
            fmt = data['formatting']
            config_dict['font_name'] = fmt.get('font_name', 'Times New Roman')
            config_dict['font_size'] = fmt.get('font_size', 12)
            config_dict['line_spacing'] = fmt.get('line_spacing', 'double')
            config_dict['margin_inches'] = fmt.get('margin_inches', 1.0)

        return cls(**config_dict)


class MarkdownToWordConverter:
    """Convert markdown legal documents to properly formatted Word documents."""

    def __init__(self, format_config: Optional[DocumentFormatConfig] = None):
        """
        Initialize converter with formatting configuration.

        Args:
            format_config: Document formatting configuration. If None, uses defaults.
        """
        self.config = format_config or DocumentFormatConfig()

    def setup_document_styles(self, doc: Document):
        """Set up proper federal court document formatting styles."""
        # Set up margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(self.config.margin_inches)
            section.bottom_margin = Inches(self.config.margin_inches)
            section.left_margin = Inches(self.config.margin_inches)
            section.right_margin = Inches(self.config.margin_inches)

        # Modify Normal style for body text
        style = doc.styles['Normal']
        font = style.font
        font.name = self.config.font_name
        font.size = Pt(self.config.font_size)

        paragraph_format = style.paragraph_format
        if self.config.line_spacing == "double":
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        else:
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

    def add_court_header(self, doc: Document):
        """Add the standard federal court header - centered and bold."""
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run(self.config.court_name)
        run.bold = True
        run.font.size = Pt(self.config.font_size)
        run.font.name = self.config.font_name

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(self.config.court_district)
        run.bold = True
        run.font.size = Pt(self.config.font_size)
        run.font.name = self.config.font_name

        if self.config.court_division:
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p3.add_run(self.config.court_division)
            run.bold = True
            run.font.size = Pt(self.config.font_size)
            run.font.name = self.config.font_name

        # Add single blank line after header
        doc.add_paragraph()

    def add_case_caption(self, doc: Document):
        """Add the case caption with perfect bracket alignment and vertical centering."""
        # Create a table for proper alignment (invisible borders)
        # Calculate number of rows needed based on parties
        num_plaintiffs = len(self.config.plaintiffs)
        num_defendants = len(self.config.defendants)

        # Rows: plaintiffs + plaintiff label + blank + v. + blank + defendants + defendant label
        total_rows = num_plaintiffs + 1 + 1 + 1 + 1 + num_defendants + 1

        table = doc.add_table(rows=total_rows, cols=3)
        table.autofit = False
        table.allow_autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths for perfect alignment
        # Column 0: Party names (left-aligned)
        # Column 1: Brackets (centered in column, 0.75" wide)
        # Column 2: Case number area (centered)
        table.columns[0].width = Inches(2.75)  # Party names
        table.columns[1].width = Inches(0.75)  # Bracket column (centered)
        table.columns[2].width = Inches(3.25)  # Case number area

        # Remove all borders
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        current_row = 0

        # Add plaintiffs
        for plaintiff in self.config.plaintiffs:
            cell_text = table.rows[current_row].cells[0].paragraphs[0]
            if plaintiff != self.config.plaintiffs[-1]:  # Not last plaintiff
                cell_text.text = f"{plaintiff.name},"
            else:
                cell_text.text = plaintiff.name
            cell_text.runs[0].font.name = self.config.font_name
            cell_text.runs[0].font.size = Pt(self.config.font_size)

            # Add bracket
            bracket = table.rows[current_row].cells[1].paragraphs[0]
            bracket.text = ")"
            bracket.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bracket.runs[0].font.name = self.config.font_name
            bracket.runs[0].font.size = Pt(self.config.font_size)

            current_row += 1

        # Add plaintiff label
        plaintiff_label = table.rows[current_row].cells[0].paragraphs[0]
        plaintiff_label.text = f"{self.config.plaintiff_label}{'s' if num_plaintiffs > 1 else ''},"
        plaintiff_label.runs[0].font.name = self.config.font_name
        plaintiff_label.runs[0].font.size = Pt(self.config.font_size)

        bracket = table.rows[current_row].cells[1].paragraphs[0]
        bracket.text = ")"
        bracket.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bracket.runs[0].font.name = self.config.font_name
        bracket.runs[0].font.size = Pt(self.config.font_size)
        current_row += 1

        # Blank line
        bracket = table.rows[current_row].cells[1].paragraphs[0]
        bracket.text = ")"
        bracket.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bracket.runs[0].font.name = self.config.font_name
        bracket.runs[0].font.size = Pt(self.config.font_size)
        current_row += 1

        # v. line - this is where we center the case number vertically
        v_line = table.rows[current_row].cells[0].paragraphs[0]
        v_line.text = "v."
        v_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        v_line.runs[0].font.name = self.config.font_name
        v_line.runs[0].font.size = Pt(self.config.font_size)

        bracket = table.rows[current_row].cells[1].paragraphs[0]
        bracket.text = ")"
        bracket.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bracket.runs[0].font.name = self.config.font_name
        bracket.runs[0].font.size = Pt(self.config.font_size)

        # Merge cells in column 2 for vertical centering of case number
        # Calculate which rows to merge (should span the middle of the caption)
        case_start_row = max(0, (total_rows // 2) - 1)
        case_end_row = min(total_rows - 1, case_start_row + 2)

        # Merge cells for case number
        merged_cell = table.rows[case_start_row].cells[2]
        for i in range(case_start_row + 1, case_end_row + 1):
            if i < total_rows:
                merged_cell.merge(table.rows[i].cells[2])

        # Add case number to merged cell with vertical centering
        case_para = merged_cell.paragraphs[0]
        case_para.text = self.config.case_number
        case_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        case_para.runs[0].bold = True
        case_para.runs[0].font.name = self.config.font_name
        case_para.runs[0].font.size = Pt(self.config.font_size)

        # Set vertical alignment for merged cell
        tc = merged_cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

        current_row += 1

        # Add defendants (skip blank line if no more rows available)
        if current_row < total_rows:
            # Blank line before defendants
            bracket = table.rows[current_row].cells[1].paragraphs[0]
            bracket.text = ")"
            bracket.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bracket.runs[0].font.name = self.config.font_name
            bracket.runs[0].font.size = Pt(self.config.font_size)
            current_row += 1

        # Add defendants
        for defendant in self.config.defendants:
            cell_text = table.rows[current_row].cells[0].paragraphs[0]
            if defendant != self.config.defendants[-1]:  # Not last defendant
                cell_text.text = f"{defendant.name},"
            else:
                cell_text.text = defendant.name
            cell_text.runs[0].font.name = self.config.font_name
            cell_text.runs[0].font.size = Pt(self.config.font_size)

            # Add bracket
            bracket = table.rows[current_row].cells[1].paragraphs[0]
            bracket.text = ")"
            bracket.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bracket.runs[0].font.name = self.config.font_name
            bracket.runs[0].font.size = Pt(self.config.font_size)

            current_row += 1

        # Add defendant label
        defendant_label = table.rows[current_row].cells[0].paragraphs[0]
        defendant_label.text = f"{self.config.defendant_label}{'s' if num_defendants > 1 else ''}."
        defendant_label.runs[0].font.name = self.config.font_name
        defendant_label.runs[0].font.size = Pt(self.config.font_size)

        bracket = table.rows[current_row].cells[1].paragraphs[0]
        bracket.text = ")"
        bracket.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bracket.runs[0].font.name = self.config.font_name
        bracket.runs[0].font.size = Pt(self.config.font_size)

        # Set uniform row height for compactness
        for row in table.rows:
            row.height = Inches(0.2)

    def add_document_title(self, doc: Document, title: str):
        """Add the document title - centered, bold, all caps."""
        # Add small space before title
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(self.config.font_size)
        run.font.name = self.config.font_name

    def process_markdown_line(self, line: str, doc: Document, in_verification: bool = False,
                               in_signature: bool = False) -> tuple[bool, bool]:
        """
        Process a single markdown line and add it to the document.

        Returns:
            Tuple of (in_verification, in_signature) flags
        """
        # Skip horizontal rules (section dividers in markdown)
        if line.strip() == '---':
            return in_verification, in_signature

        # Skip title (H1) - handled separately
        if line.startswith('# '):
            return in_verification, in_signature

        # Handle H2 headings (major sections)
        if line.startswith('## '):
            text = line[3:].strip()
            # Check if this is VERIFICATION section
            if 'VERIFICATION' in text.upper():
                in_verification = True
            # Check if this is CERTIFICATE OF SERVICE - add page break
            if 'CERTIFICATE OF SERVICE' in text.upper():
                doc.add_page_break()

            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(self.config.font_size)
            run.font.name = self.config.font_name
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)
            return in_verification, in_signature

        # Handle H3 headings (subsections)
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(self.config.font_size)
            run.font.name = self.config.font_name
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            return in_verification, in_signature

        # Handle H4 headings (sub-subsections)
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(self.config.font_size)
            run.font.name = self.config.font_name
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(0)
            return in_verification, in_signature

        # Handle blockquotes (lines starting with >)
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()  # Remove > and trim
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)  # Indent blockquote
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Process bold/italic in blockquote text
            text = re.sub(r'\*\*(.*?)\*\*', lambda m: f'{{BOLD}}{m.group(1)}{{/BOLD}}', text)
            parts = re.split(r'(\{BOLD\}|\{/BOLD\})', text)
            is_bold = False
            for part in parts:
                if part == '{BOLD}':
                    is_bold = True
                elif part == '{/BOLD}':
                    is_bold = False
                elif part:
                    run = p.add_run(part)
                    run.font.name = self.config.font_name
                    run.font.size = Pt(self.config.font_size)
                    if is_bold:
                        run.bold = True
            return in_verification, in_signature

        # Handle bullet lists (lines starting with -)
        if line.strip().startswith('- '):
            text = line.strip()[2:]  # Remove - and trim
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            # Process bold and italic markdown in bullet text
            # First, handle bold+italic (***text***)
            text = re.sub(r'\*\*\*(.*?)\*\*\*', lambda m: f'{{BOLDITALIC}}{m.group(1)}{{/BOLDITALIC}}', text)
            # Then handle bold (**text**)
            text = re.sub(r'\*\*(.*?)\*\*', lambda m: f'{{BOLD}}{m.group(1)}{{/BOLD}}', text)
            # Then handle italic (*text*)
            text = re.sub(r'\*(.*?)\*', lambda m: f'{{ITALIC}}{m.group(1)}{{/ITALIC}}', text)

            # Split by formatting markers and add runs
            parts = re.split(r'(\{BOLD\}|\{/BOLD\}|\{ITALIC\}|\{/ITALIC\}|\{BOLDITALIC\}|\{/BOLDITALIC\})', text)

            bold = False
            italic = False

            for part in parts:
                if part == '{BOLD}':
                    bold = True
                elif part == '{/BOLD}':
                    bold = False
                elif part == '{ITALIC}':
                    italic = True
                elif part == '{/ITALIC}':
                    italic = False
                elif part == '{BOLDITALIC}':
                    bold = True
                    italic = True
                elif part == '{/BOLDITALIC}':
                    bold = False
                    italic = False
                elif part:
                    run = p.add_run(part)
                    run.bold = bold
                    run.italic = italic
                    run.font.name = self.config.font_name
                    run.font.size = Pt(self.config.font_size)

            return in_verification, in_signature

        # Skip blank lines - double-spacing handles paragraph separation
        if not line.strip():
            return in_verification, in_signature

        # Handle signature lines and date lines (with more vertical space)
        if line.strip().startswith('_____') or line.strip().startswith('Dated:') or line.strip().startswith('Executed on'):
            # Add more space before signature lines
            if line.strip().startswith('_____'):
                # Add two blank lines before signature line
                doc.add_paragraph()
                doc.add_paragraph()
                # Start signature block - single spacing for attorney info
                in_signature = True
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_before = Pt(12) if line.strip().startswith('Dated:') or line.strip().startswith('Executed on') else Pt(0)
            # Signature blocks are single-spaced
            if in_signature:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            return in_verification, in_signature

        # Handle "Respectfully submitted" (with more vertical space)
        if 'Respectfully submitted' in line:
            # Add three blank lines before "Respectfully submitted"
            doc.add_paragraph()
            doc.add_paragraph()
            doc.add_paragraph()
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_before = Pt(12)
            # Start signature block after "Respectfully submitted"
            in_signature = True
            return in_verification, in_signature

        # Handle regular paragraphs
        p = doc.add_paragraph()

        # For verification section or signature blocks, use single spacing
        if in_verification or in_signature:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Process bold and italic markdown
        # First, handle bold+italic (***text***)
        line = re.sub(r'\*\*\*(.*?)\*\*\*', lambda m: f'{{BOLDITALIC}}{m.group(1)}{{/BOLDITALIC}}', line)
        # Then handle bold (**text**)
        line = re.sub(r'\*\*(.*?)\*\*', lambda m: f'{{BOLD}}{m.group(1)}{{/BOLD}}', line)
        # Then handle italic (*text*)
        line = re.sub(r'\*(.*?)\*', lambda m: f'{{ITALIC}}{m.group(1)}{{/ITALIC}}', line)

        # Split by formatting markers and add runs
        parts = re.split(r'(\{BOLD\}|\{/BOLD\}|\{ITALIC\}|\{/ITALIC\}|\{BOLDITALIC\}|\{/BOLDITALIC\})', line)

        bold = False
        italic = False

        for part in parts:
            if part == '{BOLD}':
                bold = True
            elif part == '{/BOLD}':
                bold = False
            elif part == '{ITALIC}':
                italic = True
            elif part == '{/ITALIC}':
                italic = False
            elif part == '{BOLDITALIC}':
                bold = True
                italic = True
            elif part == '{/BOLDITALIC}':
                bold = False
                italic = False
            elif part:
                run = p.add_run(part)
                run.bold = bold
                run.italic = italic
                run.font.name = self.config.font_name
                run.font.size = Pt(self.config.font_size)

        return in_verification, in_signature

    def convert(self, md_file: str, output_file: Optional[str] = None,
                skip_header: bool = True) -> str:
        """
        Convert a markdown file to a properly formatted Word document.

        Args:
            md_file: Path to markdown file
            output_file: Path for output Word file (optional, defaults to same name with .docx)
            skip_header: If True, skips markdown header section (between --- separators)

        Returns:
            Path to created Word document
        """
        if output_file is None:
            output_file = str(Path(md_file).with_suffix('.docx'))

        print(f"Converting {Path(md_file).name}...")

        # Create document
        doc = Document()

        # Set up styles
        self.setup_document_styles(doc)

        # Add court header
        self.add_court_header(doc)

        # Add case caption
        self.add_case_caption(doc)

        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # Extract title from first H1
        title = None
        for line in lines:
            if line.startswith('# '):
                title = line[2:].strip().replace('(Jury Trial Demanded)', '').strip()
                break

        # Add document title
        if title:
            self.add_document_title(doc, title)

        # Skip markdown header section if requested
        start_index = 0
        if skip_header:
            separator_count = 0
            for i, line in enumerate(lines):
                if line.strip() == '---':
                    separator_count += 1
                    if separator_count == 2:
                        start_index = i + 1
                        break

            # Also skip the H2 heading that repeats the document title
            title_upper = title.upper() if title else ""
            if start_index < len(lines):
                next_line = lines[start_index].strip()
                if next_line.startswith('## ') and title_upper in next_line.upper():
                    start_index += 1

        # Process each line starting after the header section
        in_verification = False
        in_signature = False
        for line in lines[start_index:]:
            in_verification, in_signature = self.process_markdown_line(line.rstrip('\n'), doc, in_verification, in_signature)

        # Save document
        doc.save(output_file)
        print(f"[OK] Created {Path(output_file).name}")

        return output_file


def convert_markdown_to_word(
    md_file: str,
    output_file: Optional[str] = None,
    court_config: Optional[Dict] = None
) -> str:
    """
    Convenience function to convert markdown to Word with optional court configuration.

    Args:
        md_file: Path to markdown file
        output_file: Path for output Word file (optional)
        court_config: Dictionary of court configuration values (optional)

    Returns:
        Path to created Word document

    Example:
        >>> convert_markdown_to_word(
        ...     "motion.md",
        ...     court_config={"case_number": "1:25-cv-12345"}
        ... )
    """
    config = DocumentFormatConfig(**court_config) if court_config else None
    converter = MarkdownToWordConverter(format_config=config)
    return converter.convert(md_file, output_file)
